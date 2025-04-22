import pandas as pd
import smtplib
import os
import ssl
import time
import re
import sys
import html
from getpass import getpass
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from email.mime.image import MIMEImage
from tkinter import Tk, filedialog, Text, Button, Toplevel, Entry, Label, StringVar
from docx import Document
from docx.oxml.ns import qn
from email_validator import validate_email, EmailNotValidError

# ===== CONFIGURAÇÕES =====
if getattr(sys, 'frozen', False):
    BASE_DIR = sys._MEIPASS
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

CONFIG = {
    "SMTP_SERVER": "mail.prioritabrasil.com.br",
    "SMTP_PORT": 465,
    "EMAILS_PER_SESSION": 50,
    "DELAY_BETWEEN_EMAILS": 1
}

# ===== FUNÇÕES AUXILIARES =====
def criar_conexao_smtp():
    context = ssl.create_default_context()
    context.options &= ~ssl.OP_NO_TLSv1
    context.options &= ~ssl.OP_NO_TLSv1_1
    context.minimum_version = ssl.TLSVersion.TLSv1
    try:
        server = smtplib.SMTP_SSL(CONFIG["SMTP_SERVER"], CONFIG["SMTP_PORT"], context=context)
        return server
    except ssl.SSLError as e:
        print(f"❌ Erro SSL ao conectar ao servidor SMTP: {e}")
        print("ℹ️ Isso pode indicar que o servidor não suporta a versão TLS usada. Tente usar a porta 587 com STARTTLS.")
        sys.exit(1)

def tentar_login(server, email, senha):
    try:
        server.login(email, senha)
        print("🔐 Login realizado com sucesso!")
        return True
    except smtplib.SMTPAuthenticationError:
        print("❌ Falha na autenticação: E-mail ou senha incorretos.")
        return False
    except Exception as e:
        print(f"❌ Erro ao tentar login: {e}")
        return False

def limpar_e_validar_email(email):
    email = str(email).strip()
    try:
        valid = validate_email(email, check_deliverability=False)
        return valid.email
    except EmailNotValidError:
        return None

def obter_senha():
    try:
        def salvar_senha():
            senha_var.set(entry.get())
            janela.destroy()

        janela = Toplevel()
        janela.title("Digite a Senha")
        janela.geometry("300x150")
        
        Label(janela, text="Digite a senha do e-mail:", font=("Arial", 12)).pack(pady=10)
        
        senha_var = StringVar()
        entry = Entry(janela, textvariable=senha_var, show="*", font=("Arial", 12), width=20)
        entry.pack(pady=10)
        
        Button(janela, text="Confirmar", command=salvar_senha).pack(pady=10)
        
        janela.wait_window()
        senha = senha_var.get().strip()
        if not senha:
            print("❌ Nenhuma senha fornecida. Encerrando.")
            sys.exit(1)
        return senha
    except Exception as e:
        print(f"❌ Erro ao abrir a janela de senha (tkinter): {e}")
        print("ℹ️ Digite a senha no terminal (a entrada não será exibida):")
        senha = getpass("🔑 Senha: ").strip()
        if not senha:
            print("❌ Nenhuma senha fornecida. Encerrando.")
            sys.exit(1)
        return senha

def preparar_texto(template, dados):
    placeholders = re.findall(r'\{([^}]+)\}', template)
    for placeholder in placeholders:
        if placeholder not in dados:
            raise ValueError(f"Placeholder inválido no template: '{placeholder}'")
    try:
        return template.format(**dados)
    except KeyError as e:
        raise ValueError(f"Erro ao formatar o template: {e}")

def janela_corpo_email():
    try:
        def salvar_texto():
            editor.corpo = text_widget.get("1.0", "end").strip()
            editor.destroy()

        editor = Toplevel()
        editor.title("Digite o Corpo do E-mail")
        editor.geometry("700x500")
        text_widget = Text(editor, wrap="word", font=("Arial", 12))
        text_widget.pack(padx=10, pady=10, expand=True, fill="both")
        text_widget.insert("1.0", "Use {name} para o nome e {empresa} para a empresa.\nExemplo:\nOlá {name},\nSegue proposta para {empresa}.")
        Button(editor, text="Salvar", command=salvar_texto).pack(pady=10)
        editor.wait_window()
        return editor.corpo
    except Exception as e:
        print(f"❌ Erro ao abrir a janela de edição (tkinter): {e}")
        print("ℹ️ Digite o corpo do e-mail diretamente no terminal:")
        print("Use {name} para o nome e {empresa} para a empresa.")
        print("Pressione Enter duas vezes para finalizar.")
        linhas = []
        while True:
            linha = input()
            if linha == "":
                break
            linhas.append(linha)
        return "\n".join(linhas)

def processar_docx(caminho_docx, dados):
    try:
        doc = Document(caminho_docx)
        html_parts = []

        paragraphs = list(doc.paragraphs)
        tables = list(doc.tables)
        para_index = 0
        table_index = 0

        for element in doc.element.body:
            if element.tag.endswith('}p'):
                if para_index < len(paragraphs):
                    para = paragraphs[para_index]
                    text = para.text
                    if not text:
                        html_parts.append("<tr><td><p> </p></td></tr>")
                    else:
                        formatted_text = html.escape(text).replace("\n", "<br>")
                        formatted_text = re.sub(r' {2,}', lambda m: ' ' * len(m.group(0)), formatted_text)
                        formatted_text = preparar_texto(formatted_text, dados)
                        style = para.style.name.lower()
                        if 'bold' in style:
                            html_parts.append(f"<tr><td><b>{formatted_text}</b></td></tr>")
                        elif 'italic' in style:
                            html_parts.append(f"<tr><td><i>{formatted_text}</i></td></tr>")
                        else:
                            html_parts.append(f"<tr><td>{formatted_text}</td></tr>")
                    para_index += 1
            elif element.tag.endswith('}tbl'):
                if table_index < len(tables):
                    table = tables[table_index]
                    html_parts.append("<tr><td><table border='1' cellspacing='0' cellpadding='5' width='100%' style='border-collapse: collapse;'>")
                    for row in table.rows:
                        html_parts.append("<tr>")
                        for cell in row.cells:
                            cell_text = html.escape(cell.text).replace("\n", "<br>")
                            cell_text = re.sub(r' {2,}', lambda m: ' ' * len(m.group(0)), cell_text)
                            cell_text = preparar_texto(cell_text, dados)
                            html_parts.append(f"<td>{cell_text}</td>")
                        html_parts.append("</tr>")
                    html_parts.append("</table></td></tr>")
                    table_index += 1

        corpo_template = "".join(html_parts)
        print("Debug: corpo_template gerado =", corpo_template)
        return corpo_template
    except Exception as e:
        print(f"❌ Erro ao ler o arquivo Word: {e}")
        sys.exit(1)

def ler_planilha(caminho_planilha):
    try:
        df = pd.read_excel(caminho_planilha)
        df.columns = [col.strip().title() for col in df.columns]
        required_columns = {"Name", "Email", "Company"}
        if not required_columns.issubset(df.columns):
            raise ValueError(f"Colunas obrigatórias ausentes: {required_columns}")
        return df
    except Exception as e:
        print(f"❌ Erro ao ler a planilha: {e}")
        sys.exit(1)

def preparar_email(assunto_template, corpo_template, row, caminhos_anexos, assinatura_imagem, remetente_email):
    dados = {"name": row["Name"], "empresa": row["Company"]}
    assunto = preparar_texto(assunto_template, dados)

    msg = MIMEMultipart("related")
    msg["From"] = remetente_email
    msg["To"] = row["Email"]
    msg["Subject"] = assunto

    corpo_html = f"""
    <html>
    <head>
        <style>
            table.email-container {{ width: 100%; max-width: 600px; margin: 0 auto; }}
            p {{ margin: 0.5em 0; }}
        </style>
    </head>
    <body>
        <table class="email-container" width="100%" cellpadding="0" cellspacing="0">
            {corpo_template}
            <tr><td><br><br></td></tr>
            <tr><td><img src="cid:assinatura" width="400"></td></tr>
        </table>
    </body>
    </html>
    """

    msg_alt = MIMEMultipart("alternative")
    msg.attach(msg_alt)
    msg_alt.attach(MIMEText(corpo_html, "html"))

    for caminho in caminhos_anexos:
        if not os.path.exists(caminho):
            print(f"⚠️ Arquivo não encontrado: {caminho}")
            continue
        with open(caminho, "rb") as f:
            anexo = MIMEApplication(f.read(), Name=os.path.basename(caminho))
            anexo["Content-Disposition"] = f'attachment; filename="{os.path.basename(caminho)}"'
            msg.attach(anexo)

    if assinatura_imagem:
        if os.path.exists(assinatura_imagem):
            try:
                with open(assinatura_imagem, "rb") as img:
                    imagem = MIMEImage(img.read())
                    imagem.add_header("Content-ID", "<assinatura>")
                    imagem.add_header("Content-Disposition", "inline", filename=os.path.basename(assinatura_imagem))
                    msg.attach(imagem)
            except Exception as e:
                print(f"⚠️ Erro ao anexar assinatura: {e}")
        else:
            print("⚠️ Arquivo de assinatura não encontrado. Continuando sem imagem.")
    else:
        print("ℹ️ Nenhuma imagem de assinatura selecionada. Continuando sem imagem.")

    return msg

def enviar_email(server, msg, log_file):
    try:
        server.send_message(msg)
        print(f"✅ E-mail enviado para: {msg['To']}")
        return True
    except smtplib.SMTPRecipientsRefused as e:
        print(f"❌ E-mail inválido ou recusado: {msg['To']}. Detalhes: {e}")
        log_file.write(f"[FALHA] {msg['To']}: E-mail inválido ou recusado - {e}\n")
        return False
    except smtplib.SMTPServerDisconnected as e:
        print(f"❌ Conexão com o servidor SMTP perdida: {e}")
        log_file.write(f"[FALHA] {msg['To']}: Conexão perdida - {e}\n")
        return False
    except smtplib.SMTPException as e:
        print(f"❌ Falha ao enviar para {msg['To']}: {e}")
        log_file.write(f"[FALHA] {msg['To']}: Falha no envio - {e}\n")
        return False
    except Exception as e:
        print(f"❌ Erro inesperado ao enviar para {msg['To']}: {e}")
        log_file.write(f"[FALHA] {msg['To']}: Erro inesperado - {e}\n")
        return False

# ===== FLUXO PRINCIPAL =====
def main():
    try:
        root = Tk()
        root.withdraw()  # Esconde a janela principal do Tkinter
    except Exception as e:
        print(f"❌ Erro ao inicializar o Tkinter: {e}")
        print("ℹ️ Seleção de arquivos será feita via entrada manual no terminal.")

    print("📧 Digite o e-mail que será usado para enviar os e-mails:")
    remetente_email = input("> ").strip()
    if limpar_e_validar_email(remetente_email) is None:
        print("❌ E-mail remetente inválido. Encerrando.")
        sys.exit(1)

    # Obter a senha usando janela segura ou getpass
    senha = obter_senha()

    print("Selecione a planilha de contatos (.xlsx)")
    try:
        caminho_planilha = filedialog.askopenfilename(filetypes=[("Planilhas Excel", "*.xlsx")])
    except Exception as e:
        print(f"❌ Erro ao abrir a janela de seleção de arquivos: {e}")
        print("ℹ️ Digite o caminho completo para a planilha (.xlsx):")
        caminho_planilha = input("> ").strip()
    if not caminho_planilha:
        print("❌ Nenhuma planilha selecionada. Encerrando.")
        sys.exit(1)

    df = ler_planilha(caminho_planilha)

    print("\nDigite o assunto do e-mail (use {name} e {empresa}):")
    assunto_template = input("> ").strip()

    print("\nComo você deseja incluir o corpo do e-mail?")
    print("1 - Digitar manualmente (janela grande)")
    print("2 - Importar de um arquivo .docx (Word)")
    escolha = input("Digite 1 ou 2: ").strip()

    if escolha == "1":
        corpo_template = janela_corpo_email()
    elif escolha == "2":
        print("Selecione o arquivo .docx com o corpo do e-mail...")
        try:
            caminho_docx = filedialog.askopenfilename(filetypes=[("Documentos Word", "*.docx")])
        except Exception as e:
            print(f"❌ Erro ao abrir a janela de seleção de arquivos: {e}")
            print("ℹ️ Digite o caminho completo para o arquivo .docx:")
            caminho_docx = input("> ").strip()
        if not caminho_docx:
            print("❌ Nenhum arquivo .docx selecionado.")
            sys.exit(1)
    else:
        print("❌ Opção inválida.")
        sys.exit(1)

    if escolha == "1" and (not corpo_template or corpo_template.strip() == ""):
        print("❌ Corpo do e-mail não foi informado. Encerrando.")
        sys.exit(1)

    tem_anexo = input("Deseja adicionar anexos? (s/n): ").strip().lower()
    caminhos_anexos = []
    if tem_anexo == "s":
        print("Selecione um ou mais arquivos a serem anexados...")
        try:
            caminhos_anexos = filedialog.askopenfilenames()
        except Exception as e:
            print(f"❌ Erro ao abrir a janela de seleção de arquivos: {e}")
            print("ℹ️ Digite os caminhos completos dos arquivos (um por linha, pressione Enter duas vezes para finalizar):")
            while True:
                caminho = input("> ").strip()
                if caminho == "":
                    break
                caminhos_anexos.append(caminho)
        if not caminhos_anexos:
            print("Nenhum arquivo selecionado. Continuando sem anexo.")

    tem_assinatura = input("Deseja adicionar uma imagem de assinatura ao e-mail? (s/n): ").strip().lower()
    assinatura_imagem = None
    if tem_assinatura == "s":
        print("Selecione a imagem de assinatura (.png, .jpg, etc.)...")
        try:
            assinatura_imagem = filedialog.askopenfilename(filetypes=[("Imagens", "*.png *.jpg *.jpeg *.gif")])
        except Exception as e:
            print(f"❌ Erro ao abrir a janela de seleção de arquivos: {e}")
            print("ℹ️ Digite o caminho completo para a imagem de assinatura:")
            assinatura_imagem = input("> ").strip()
        if not assinatura_imagem:
            print("Nenhuma imagem de assinatura selecionada. Continuando sem imagem.")

    with open("log_envio.txt", "a", encoding="utf-8") as log_file:
        server = criar_conexao_smtp()
        if not tentar_login(server, remetente_email, senha):
            try:
                server.quit()
            except:
                pass
            sys.exit(1)

        # Limpar a senha da memória após o login
        senha = None

        emails_enviados = 0
        total_emails = len(df)

        for idx, row in df.iterrows():
            email = limpar_e_validar_email(row["Email"])
            if not email:
                print(f"⚠️ E-mail inválido: '{row['Email']}'. Pulando.")
                log_file.write(f"[ERRO] E-mail inválido: {row['Email']}\n")
                continue

            row["Email"] = email
            dados = {"name": row["Name"], "empresa": row["Company"]}
            if escolha == "2":
                corpo_template = processar_docx(caminho_docx, dados)
            if not corpo_template or corpo_template.strip() == "":
                print(f"❌ Corpo do e-mail vazio para {email}. Pulando.")
                log_file.write(f"[ERRO] Corpo do e-mail vazio: {email}\n")
                continue

            msg = preparar_email(assunto_template, corpo_template, row, caminhos_anexos, assinatura_imagem, remetente_email)

            try:
                server.noop()
            except (smtplib.SMTPServerDisconnected, smtplib.SMTPException, ssl.SSLError) as e:
                print(f"⚠️ Conexão perdida: {e}. Reconectando...")
                try:
                    server.quit()
                except:
                    pass
                server = criar_conexao_smtp()
                # Solicitar a senha novamente se a conexão for perdida
                senha = obter_senha()
                if not tentar_login(server, remetente_email, senha):
                    try:
                        server.quit()
                    except:
                        pass
                    sys.exit(1)
                # Limpar a senha da memória novamente
                senha = None

            if enviar_email(server, msg, log_file):
                log_file.write(f"[OK] {row['Name']} - {email} - Assunto: {msg['Subject']}\n")
                emails_enviados += 1
            else:
                log_file.write(f"[FALHA] {email}: Falha no envio\n")

            if idx < total_emails - 1:
                time.sleep(CONFIG["DELAY_BETWEEN_EMAILS"])

            if emails_enviados > 0 and emails_enviados % CONFIG["EMAILS_PER_SESSION"] == 0:
                print(f"ℹ️ Enviados {emails_enviados} e-mails. Reconectando ao servidor SMTP...")
                try:
                    server.quit()
                except:
                    pass
                server = criar_conexao_smtp()
                # Solicitar a senha novamente para a nova conexão
                senha = obter_senha()
                if not tentar_login(server, remetente_email, senha):
                    try:
                        server.quit()
                    except:
                        pass
                    sys.exit(1)
                # Limpar a senha da memória novamente
                senha = None

        try:
            server.quit()
        except:
            pass

    print(f"\n✅ Processo finalizado. E-mails enviados com sucesso: {emails_enviados}/{total_emails}")

if __name__ == "__main__":
    main()